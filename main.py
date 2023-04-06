import time
from docx import Document
from docx.shared import Pt  # 用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  # 设置字体
from docx.shared import RGBColor  # 设置字体的颜色
from win32com import client as wc
import os
import re
import shutil
import zipfile


def docx_remove_content(doc, finish_doc):
    # 定义需要去除的内容
    content_to_remove = '''【提示】本文档来源自第一范文网（https://www.51xiazai.cn），第一范文网是中国最大的范文网站，专注于提供各种优质工作学习办公文档，欢迎访问。
微信搜索公众号“第一范文网”，关注后可方便查找下载各种文档。
转发文档请保留以上标记，谢谢！'''
    # 打开doc文件
    document = Document(doc)
    # 遍历doc文件中的段落
    for para in document.paragraphs:
        # 如果段落中包含需要去除的内容，使用正则表达式替换为空字符串
        if re.search(content_to_remove, para.text):
            para.text = re.sub(content_to_remove, '', para.text)
        # else:
        #     doc_basename = os.path.basename(doc).split("】")[1]
        #     doc_title_name = os.path.splitext(doc_basename)[0]
        #
        #     # 更改段落字体颜色，例如：爱校之星班主任评语
        #     if doc_title_name == para.text:
        #         for run in para.runs:
        #             run.font.bold = True
        #             run.font.underline = True
        #             run.font.size = Pt(20)
        #             run.font.color.rgb = RGBColor(255, 0, 0)
        #         continue
        #
        #     # 删除带有篇数量的段落，例如：爱校之星班主任评语（精选25篇）
        #     if doc_title_name + "（" in para.text:
        #         p = para._element
        #         p.getparent().remove(p)
        #         para._p = para._element = None
        #         continue
        #
        #     # 更改段落字体颜色，例如：爱校之星班主任评语 篇1
        #     if doc_title_name + " 篇" in para.text:
        #         for run in para.runs:
        #             run.font.bold = True
        #             run.font.size = Pt(15)
        #             run.font.color.rgb = RGBColor(255, 0, 0)

    document.save(finish_doc)


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    # finish_doc： 需要另存为的新文件名
    document = Document(doc)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    document.save(doc)


def change_word_font(doc):
    # 打开doc文件
    document = Document(doc)
    document.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
    document.save(doc)


def get_word_pages(in_file):
    pages = 1
    try:
        word = wc.Dispatch("Word.Application")
        try:
            doc = word.Documents.Open(in_file)
            word.ActiveDocument.Repaginate()
            pages = word.ActiveDocument.ComputeStatistics(2)
            doc.Close()
            word.Quit()
            return pages
        except Exception as e:
            print(e)
        finally:
            return pages
    except Exception as e:
        print(e)
    finally:
        return pages


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            print('转换成功')
            doc.Close()
            word.Quit()
        except Exception as e:
            print(1111)
            print(e)
    except Exception as e:
        print(e)


def support_gbk(zip_file: zipfile):
    name_to_info = zip_file.NameToInfo
    # copy map first
    for name, info in name_to_info.copy().items():
        real_name = name.encode('cp437').decode('gbk')
        if real_name != name:
            info.filename = real_name
            del name_to_info[name]
            name_to_info[real_name] = info
    return zip_file


AllCategory = ["word模板"]

if __name__ == '__main__':
    # for category in AllCategory:
    #     category_path = "G:\\\www.ppt818.com\\" + category
    #     child_category_zips = sorted(os.listdir(category_path))
    #     for zipFile in child_category_zips:
    #         if ".zip" in zipFile:
    #             with support_gbk(zipfile.ZipFile(category_path + "\\" + zipFile, 'r')) as zf:  # 压缩文件位置
    #                 for file in zf.namelist():
    #                     if ".doc" in file:
    #                         print(file)
    #                         file_size = zf.getinfo(file).file_size
    #                         new_path = "G:\\www.ppt818.com\\" + category + "\\" + zipFile.split(".")[0] + "." + file.split(".")[1]
    #                         if file_size > 0:
    #                             # 是文件，通过open创建文件，写入数据
    #                             with open(new_path, 'wb') as f:
    #                                 # zf.read 是读取压缩包里的文件内容
    #                                 f.write(zf.read(file))
    #
    # exit(1)

    # for category in AllCategory:
    #     category_path = "G:\\www.pc6.com\\" + category
    #     files = sorted(os.listdir(category_path))
    #     for file in files:
    #         print(category_path + "\\" + file)
    #         os.rename(category_path + "\\" + file, "G:\\www.pc6.com\\" + "【" + category + "】" + file)
    # for child_category in child_category_dirs:
    #     word_dir = category_path + "\\" + child_category
    #
    #     files = sorted(os.listdir(child_category_dirs))
    #     for file in files:
    #         print(word_dir + "\\" + file)
    #         exit(1)
    #         os.rename(word_dir + "\\" + file, word_dir + "\\" + "【"+category+"】" + file)

    for category in AllCategory:
        category_path = "G:\\www.ppt818.com\\" + category
        files = sorted(os.listdir(category_path))
        for file in files:
            if os.path.splitext(file)[1] in [".doc", ".docx"]:
                print(file)

                finish_file = category_path + "\\" + os.path.splitext(file)[0] + ".docx"
                word_file = category_path + "\\" + file

                if not os.path.exists(finish_file):
                    if os.path.splitext(file)[1] == ".doc":
                        # 将doc文件转化为docx文件
                        print("转化文件")
                        doc2docx(word_file, finish_file)
                        os.remove(word_file)

                # 删除页眉页脚
                remove_header_footer(finish_file)
                # 改变文档字体
                change_word_font(finish_file)
